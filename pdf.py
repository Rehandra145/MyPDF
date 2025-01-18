import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from PyPDF2 import PdfReader, PdfWriter
from PIL import Image, ImageTk
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import customtkinter as ctk
from typing import List, Optional

class ModernPDFTool:
    def __init__(self):
        self.root = ctk.CTk()
        self.root.title("PDF Manager")
        self.root.geometry("1000x700")
        ctk.set_appearance_mode("light")
        ctk.set_default_color_theme("blue")
        
        self.current_mode = None
        self.image_files: List[str] = []
        self.current_images: List[Image.Image] = []
        self.current_image_index = 0
        self.image_order: List[int] = []
        
        self.filename_label = ttk.Label(
            self.root,
            text="No PDF selected",
            font=("Helvetica", 10)
        )
        self.filename_label.pack(pady=5)
        
        self.setup_ui()
        
    def setup_ui(self):
        # Main container
        self.main_frame = ctk.CTkFrame(self.root)
        self.main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Left panel for controls
        self.left_panel = ctk.CTkFrame(self.main_frame, width=250)
        self.left_panel.pack(side="left", fill="y", padx=(0, 10))
        self.left_panel.pack_propagate(False)  # Prevent shrinking
        
        # Right panel for preview
        self.right_panel = ctk.CTkFrame(self.main_frame)
        self.right_panel.pack(side="right", fill="both", expand=True)
        
        self.setup_left_panel()
        self.setup_right_panel()
        
    def setup_left_panel(self):
        # Title
        ctk.CTkLabel(
            self.left_panel,
            text="PDF Manager",
            font=("Helvetica", 20, "bold")
        ).pack(pady=(20, 30))
        
        # Conversion buttons
        conversion_frame = ctk.CTkFrame(self.left_panel)
        conversion_frame.pack(fill="x", padx=20)
        
        buttons = [
            ("Word → PDF", lambda: self.start_conversion("word")),
            ("Images → PDF", lambda: self.start_conversion("images")),
            ("Crop PDF", lambda: self.start_conversion("crop"))
        ]
        
        for text, command in buttons:
            ctk.CTkButton(
                conversion_frame,
                text=text,
                command=command,
                width=200,
                height=40,
                corner_radius=8
            ).pack(pady=10)
        
        # Crop controls (hidden by default)
        self.crop_frame = ctk.CTkFrame(self.left_panel)
        
        ctk.CTkLabel(self.crop_frame, text="Page Range").pack(pady=(10, 5))
        range_frame = ctk.CTkFrame(self.crop_frame)
        range_frame.pack(fill="x", padx=10)
        
        ctk.CTkLabel(range_frame, text="Start:").pack(side="left", padx=5)
        self.crop_start = ctk.CTkEntry(range_frame, width=50)
        self.crop_start.pack(side="left", padx=5)
        
        ctk.CTkLabel(range_frame, text="End:").pack(side="left", padx=5)
        self.crop_end = ctk.CTkEntry(range_frame, width=50)
        self.crop_end.pack(side="left", padx=5)
        
        # Image order controls (hidden by default)
        self.image_controls_frame = ctk.CTkFrame(self.left_panel)
        
        ctk.CTkButton(
            self.image_controls_frame,
            text="↑ Move Up",
            command=self.move_image_up,
            width=95
        ).pack(side="left", padx=2)
        
        ctk.CTkButton(
            self.image_controls_frame,
            text="↓ Move Down",
            command=self.move_image_down,
            width=95
        ).pack(side="left", padx=2)

    def setup_right_panel(self):
        # Preview label (hidden by default)
        self.preview_label = ctk.CTkLabel(
            self.right_panel,
            text="Preview",
            font=("Helvetica", 16, "bold")
        )
        
        # Preview canvas
        self.preview_frame = ctk.CTkFrame(self.right_panel)
        self.preview_canvas = ctk.CTkCanvas(
            self.preview_frame,
            width=600,
            height=500,
            bg="white"
        )
        self.preview_canvas.pack(pady=10, padx=10, fill="both", expand=True)
        
        # Navigation and convert frame
        self.navigation_frame = ctk.CTkFrame(self.right_panel)
        
        # Previous button
        self.prev_btn = ctk.CTkButton(
            self.navigation_frame,
            text="← Previous",
            command=self.prev_page,
            width=100
        )
        self.prev_btn.pack(side="left", padx=5)
        
        # Convert button (in the middle)
        self.convert_btn = ctk.CTkButton(
            self.navigation_frame,
            text="Convert",
            command=self.convert_files,
            width=150,
            height=40,
            fg_color="#28a745",  # Green color
            hover_color="#218838"  # Darker green for hover
        )
        self.convert_btn.pack(side="left", padx=20)
        
        # Page label
        self.page_label = ctk.CTkLabel(self.navigation_frame, text="")
        self.page_label.pack(side="left", padx=5)
        
        # Next button
        self.next_btn = ctk.CTkButton(
            self.navigation_frame,
            text="Next →",
            command=self.next_page,
            width=100
        )
        self.next_btn.pack(side="right", padx=5)

    def show_preview_elements(self):
        self.preview_label.pack(pady=(15, 5))
        self.filename_label = ctk.CTkLabel(
            self.right_panel,
            text="",
            font=("Helvetica", 12)
        )
        self.filename_label.pack(pady=(0, 5))
        self.preview_frame.pack(fill="both", expand=True, padx=10, pady=5)
        self.navigation_frame.pack(fill="x", pady=10, padx=10)

    def hide_preview_elements(self):
        self.preview_label.pack_forget()
        self.filename_label.pack_forget()
        self.preview_frame.pack_forget()
        self.navigation_frame.pack_forget()

    def start_conversion(self, mode):
        self.current_mode = mode
        self.crop_frame.pack_forget()
        self.image_controls_frame.pack_forget()
        self.hide_preview_elements()
        
        if mode == "word":
            self.select_word_file()
        elif mode == "images":
            self.select_images()
        elif mode == "crop":
            self.select_pdf_for_crop()

    def select_word_file(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Word Files", "*.docx")]
        )
        if file_path:
            self.current_file = file_path
            self.show_preview_elements()
            self.preview_canvas.delete("all")
            
            try:
                from docx import Document
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                
                # Calculate text height to fit canvas
                canvas_width = 580
                canvas_height = 480
                font_size = 10
                lines = text.split('\n')
                max_lines = canvas_height // (font_size + 2)  # Approximate line height
                
                pages = [lines[i:i + max_lines] for i in range(0, len(lines), max_lines)]
                
                self.current_pages = ["\n".join(page) for page in pages]
                self.current_page_index = 0
                
                self.display_current_page()
            except ImportError:
                messagebox.showerror(
                    "Error",
                    "Please install 'python-docx' package first."
                )
            except Exception as e:
                messagebox.showerror(
                    "Error",
                    f"Error reading Word file: {str(e)}"
                )

    def display_current_page(self):
        self.preview_canvas.delete("all")
        if hasattr(self, 'current_pages'):
            display_text = "\n".join(self.current_pages[self.current_page_index])
            self.preview_canvas.create_text(
                10, 10,
                text=display_text,
                width=580 - 20,
                font=("Helvetica", 10),
                anchor="nw"
            )
        self.update_page_label()

    def select_images(self):
        self.image_files = list(filedialog.askopenfilenames(
            filetypes=[("Image Files", "*.jpg *.jpeg *.png")]
        ))
        if self.image_files:
            self.current_images = []
            self.image_order = list(range(len(self.image_files)))
            self.current_image_index = 0
            
            for img_path in self.image_files:
                img = Image.open(img_path)
                self.current_images.append(img)
            
            self.show_preview_elements()
            self.image_controls_frame.pack(pady=5)
            self.display_current_image()
            self.update_page_label()

    def select_pdf_for_crop(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("PDF Files", "*.pdf")]
        )
        if file_path:
            try:
                self.current_file = file_path
                self.current_pdf = PdfReader(file_path)
                self.current_page = 0
                self.total_pages = len(self.current_pdf.pages)
                self.show_preview_elements()
                # Update filename label
                self.filename_label.configure(text=f"File: {file_path.split('/')[-1]}")
                self.crop_frame.pack(pady=5)
                self.update_preview()
            
            except Exception as e:
                messagebox.showerror(
                    "Error",
                    f"Error opening PDF: {str(e)}"
                )

    def move_image_up(self):
        if self.current_image_index > 0:
            # Swap positions in image_order
            idx = self.current_image_index
            self.image_order[idx], self.image_order[idx-1] = \
                self.image_order[idx-1], self.image_order[idx]
            self.current_image_index -= 1
            self.display_current_image()
            self.update_page_label()

    def move_image_down(self):
        if self.current_image_index < len(self.image_order) - 1:
            # Swap positions in image_order
            idx = self.current_image_index
            self.image_order[idx], self.image_order[idx+1] = \
                self.image_order[idx+1], self.image_order[idx]
            self.current_image_index += 1
            self.display_current_image()
            self.update_page_label()

    def display_current_image(self):
        self.preview_canvas.delete("all")
        if self.current_images:
            img = self.current_images[self.image_order[self.current_image_index]]
            # Calculate scaling to fit canvas while maintaining aspect ratio
            canvas_width = 580
            canvas_height = 480
            img_width, img_height = img.size
            
            # Calculate scaling factor
            scale = min(canvas_width/img_width, canvas_height/img_height)
            new_width = int(img_width * scale)
            new_height = int(img_height * scale)
            
            img = img.resize((new_width, new_height), Image.LANCZOS)
            photo = ImageTk.PhotoImage(img)
            
            # Center the image
            x = (canvas_width - new_width) // 2 + 10
            y = (canvas_height - new_height) // 2 + 10
            
            self.preview_canvas.create_image(
                x, y,
                image=photo,
                anchor="nw"
            )
            self.preview_canvas.image = photo

    def update_preview(self):
      if hasattr(self, 'current_pdf'):
        self.preview_canvas.delete("all")
        try:
            # Get current page safely
            page = self.current_pdf.pages[self.current_page]
            
            # Extract text with error handling
            try:
                text = page.extract_text()
                print(f"Extracted text: {text[:100]}...") # Debug logging
            except Exception as e:
                text = f"Error extracting text: {str(e)}"
                print(f"Text extraction error: {e}")
            
            # Display text first
            self.preview_canvas.create_text(
                10, 10,
                text=text,
                width=580,
                font=("Helvetica", 10),
                anchor="nw"
            )
            
            # Extract and display images
            canvas_width = 580
            canvas_height = 480
            
            try:
                resources = page.get('/Resources', {})
                if resources and '/XObject' in resources:
                    xobject = resources['/XObject']
                    if hasattr(xobject, 'get_object'):
                        xobject = xobject.get_object()
                    
                    print(f"Found XObjects: {list(xobject.keys())}") # Debug logging
                    
                    for key, obj in xobject.items():
                        if isinstance(obj, dict) and obj.get('/Subtype') == '/Image':
                            try:
                                from PIL import Image
                                import io
                                
                                print(f"Processing image: {key}") # Debug logging
                                
                                # Get image data
                                if obj.get('/Filter') == '/DCTDecode':
                                    img_data = obj.get_data()
                                    img = Image.open(io.BytesIO(img_data))
                                elif obj.get('/Filter') == '/FlateDecode':
                                    width = obj.get('/Width', 0)
                                    height = obj.get('/Height', 0)
                                    if width and height:
                                        img_data = obj.get_data()
                                        img = Image.frombytes('RGB', (width, height), img_data)
                                
                                if img:
                                    # Scale image
                                    img_width, img_height = img.size
                                    scale = min(canvas_width/img_width, canvas_height/img_height)
                                    new_width = int(img_width * scale)
                                    new_height = int(img_height * scale)
                                    
                                    img = img.resize((new_width, new_height), Image.LANCZOS)
                                    photo = ImageTk.PhotoImage(img)
                                    
                                    # Center image
                                    x = (canvas_width - new_width) // 2
                                    y = (canvas_height - new_height) // 2
                                    
                                    self.preview_canvas.create_image(
                                        x, y,
                                        image=photo,
                                        anchor="nw"
                                    )
                                    self.preview_canvas.image = photo
                                    print(f"Image displayed: {key}") # Debug logging
                                    
                            except Exception as e:
                                print(f"Error processing image {key}: {e}")
                                
            except Exception as e:
                print(f"Error processing resources: {e}")
            
            # Update navigation
            self.page_label.configure(
                text=f"Page {self.current_page + 1}/{self.total_pages}"
            )
            
            # Update navigation buttons
            self.prev_btn.pack_forget() if self.current_page == 0 else self.prev_btn.pack(side="left", padx=5)
            self.next_btn.pack_forget() if self.current_page >= self.total_pages - 1 else self.next_btn.pack(side="right", padx=5)
            
        except Exception as e:
            print(f"Error in update_preview: {e}")
            messagebox.showerror(
                "Error",
                f"Error displaying page: {str(e)}"
            )

    def prev_page(self):
      if self.current_mode == "crop":
        if self.current_page > 0:
            self.current_page -= 1
            self.update_preview()
      elif self.current_mode == "word":
        if self.current_page_index > 0:
            self.current_page_index -= 1
            self.display_current_page()
      elif self.current_mode == "images":
        if self.current_image_index > 0:
            self.current_image_index -= 1
            self.display_current_image()
            self.update_page_label()

    def next_page(self):
      if self.current_mode == "crop":
        if self.current_page < self.total_pages - 1:
            self.current_page += 1
            self.update_preview()
      elif self.current_mode == "word":
        if self.current_page_index < len(self.current_pages) - 1:
            self.current_page_index += 1
            self.display_current_page()
      elif self.current_mode == "images":
        if self.current_image_index < len(self.current_images) - 1:
            self.current_image_index += 1
            self.display_current_image()
            self.update_page_label()

    def update_page_label(self):
        if self.current_mode == "word" or self.current_mode == "crop":
            self.page_label.configure(
                text=f"Page {self.current_page_index + 1}/{len(self.current_pages)}"
            )
        elif self.current_mode == "images":
            self.page_label.configure(
                text=f"Image {self.current_image_index + 1}/{len(self.current_images)}"
            )
            # Show/hide navigation buttons based on current position
            if self.current_image_index == 0:
                self.prev_btn.pack_forget()
            else:
                self.prev_btn.pack(side="left", padx=5)
                
            if self.current_image_index == len(self.current_images) - 1:
                self.next_btn.pack_forget()
            else:
                self.next_btn.pack(side="right", padx=5)
                
        elif hasattr(self, 'current_pdf'):
            self.page_label.configure(
                text=f"Page {self.current_page + 1}/{self.total_pages}"
            )
            # Show/hide navigation buttons based on current position
            if self.current_page == 0:
                self.prev_btn.pack_forget()
            else:
                self.prev_btn.pack(side="left", padx=5)
                
            if self.current_page == self.total_pages - 1:
                self.next_btn.pack_forget()
            else:
                self.next_btn.pack(side="right", padx=5)

    def convert_files(self):
        if self.current_mode == "word":
            self.convert_word_to_pdf()
        elif self.current_mode == "images":
            self.convert_images_to_pdf()
        elif self.current_mode == "crop":
            self.crop_pdf_pages()

    def convert_word_to_pdf(self):
      try:
        from docx import Document
        output_file = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF Files", "*.pdf")]
        )
        if output_file:
            doc = Document(self.current_file)
            c = canvas.Canvas(output_file, pagesize=(595, 842))  # A4 size
            width, height = 595, 842
            y = height - 40  # Start from top of the page
            for para in doc.paragraphs:
                if y < 40:  # If near bottom, create new page
                    c.showPage()
                    y = height - 40
                c.drawString(40, y, para.text)
                y -= 20  # Move down for next line
            c.save()
            messagebox.showinfo(
                "Success",
                "Word file converted to PDF successfully!"
            )
      except ImportError:
        messagebox.showerror(
            "Error",
            "Please install 'python-docx' package first."
        )
      except Exception as e:
        messagebox.showerror(
            "Error",
            f"Error converting Word to PDF: {str(e)}"
        )

    def convert_images_to_pdf(self):
      if self.image_files:
        output_file = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF Files", "*.pdf")]
        )
        if output_file:
            try:
                ordered_images = [self.current_images[i].convert("RGB") 
                                  for i in self.image_order]
                a4_width, a4_height = 595, 842
                resized_images = []
                for img in ordered_images:
                    img.thumbnail((a4_width, a4_height), Image.LANCZOS)
                    new_img = Image.new("RGB", (a4_width, a4_height), (255, 255, 255))
                    new_img.paste(img, ((a4_width - img.width) // 2, (a4_height - img.height) // 2))
                    resized_images.append(new_img)
                resized_images[0].save(
                    output_file,
                    save_all=True,
                    append_images=resized_images[1:],
                    format="PDF"
                )
                messagebox.showinfo(
                    "Success",
                    "Images converted to PDF successfully!"
                )
            except Exception as e:
                messagebox.showerror(
                    "Error",
                    f"Error converting images to PDF: {str(e)}"
                )

    def crop_pdf_pages(self):
        try:
            start_page = int(self.crop_start.get()) - 1
            end_page = int(self.crop_end.get()) - 1
            if start_page < 0 or end_page >= self.total_pages or start_page > end_page:
                raise ValueError("Invalid page range")
            
            output_file = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF Files", "*.pdf")]
            )
            if output_file:
                writer = PdfWriter()
                for i in range(start_page, end_page + 1):
                    writer.add_page(self.current_pdf.pages[i])
                with open(output_file, "wb") as f:
                    writer.write(f)
                messagebox.showinfo(
                    "Success",
                    "PDF cropped successfully!"
                )
        except ValueError as e:
            messagebox.showerror(
                "Error",
                f"Invalid page range: {str(e)}"
            )
        except Exception as e:
            messagebox.showerror(
                "Error",
                f"Error cropping PDF: {str(e)}"
            )

    def run(self):
        self.root.mainloop()

if __name__ == "__main__":
    app = ModernPDFTool()
    app.run()